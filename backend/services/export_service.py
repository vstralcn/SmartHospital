from __future__ import annotations

import re
from datetime import datetime
from html.parser import HTMLParser
from pathlib import Path

from utils.file_utils import ensure_directory

_NOISE_LINES = {
    "p, li { white-space: pre-wrap; }",
    "hr { height: 1px; border-width: 0; }",
    'li.unchecked::marker { content: "\\2610"; }',
    'li.checked::marker { content: "\\2612"; }',
}


def _is_noise_text(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    if re.match(r"^源音频[:：].*", stripped):
        return True
    if re.fullmatch(r"\[Pasted ~\d+ lines\]", stripped):
        return True
    return stripped in _NOISE_LINES


class ExportService:
    def __init__(self, output_dir: Path) -> None:
        self.output_dir = output_dir

    def build_default_filename(self) -> str:
        return f"medical-emr-demo-{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx"

    def export_docx(self, payload, export_path: Path | None = None) -> Path:
        try:
            import docx
            from docx.shared import Pt as pt
            Document = docx.Document
        except ImportError as exc:
            raise RuntimeError("python-docx 未安装，暂时无法导出 Word。") from exc

        ensure_directory(self.output_dir)
        output_path = export_path or (self.output_dir / self.build_default_filename())
        document = Document()
        self._configure_document(document, pt)
        document.add_heading("医疗问诊助手 Demo 导出", level=1)
        document.add_heading("结构化字段", level=2)
        structured = payload.structured
        self._add_paragraph(document, f"主诉：{structured.chief_complaint}")
        self._add_paragraph(document, f"现病史：{structured.present_illness}")
        self._add_paragraph(document, f"既往史：{structured.past_history}")
        self._add_paragraph(document, f"手术史：{structured.surgical_history}")
        self._add_paragraph(document, f"过敏史：{structured.allergy_history}")
        self._add_paragraph(document, f"用药史：{structured.medication_history}")
        self._add_paragraph(
            document, f"待确认：{'；'.join(structured.needs_confirmation) or '无'}"
        )
        document.add_heading("病历草稿", level=2)
        sanitized_html = self._sanitize_html(getattr(payload, "emr_html", ""))
        if sanitized_html:
            self._append_rich_text(document, sanitized_html)
        else:
            self._append_plain_text(document, payload.emr_text)
        document.add_heading("风险提示", level=2)
        for item in payload.risk_alerts or ["无"]:
            self._add_paragraph(document, item, style="List Bullet")
        document.save(output_path)
        self.export_json(payload, output_path.with_suffix(".json"))
        return output_path

    def export_json(self, payload, export_path: Path | None = None) -> Path:
        ensure_directory(self.output_dir)
        output_path = export_path or (
            self.output_dir / self.build_default_filename().replace(".docx", ".json")
        )
        output_path.write_text(payload.model_dump_json(indent=2), encoding="utf-8")
        return output_path

    def _append_rich_text(self, document, html_text: str) -> None:
        parser = _DocxHTMLParser(document)
        parser.feed(html_text)
        parser.close()

    def _append_plain_text(self, document, text: str) -> None:
        cleaned_text = self._sanitize_plain_text(text)
        if not cleaned_text:
            self._add_paragraph(document, "无")
            return

        for block in cleaned_text.split("\n\n"):
            lines = [line.strip() for line in block.splitlines() if line.strip()]
            if not lines:
                continue
            paragraph = self._add_paragraph(document, lines[0])
            for line in lines[1:]:
                paragraph.add_run().add_break()
                paragraph.add_run(line)

    def _configure_document(self, document, pt) -> None:
        normal_style = document.styles["Normal"]
        normal_format = normal_style.paragraph_format
        normal_format.line_spacing = 1.25
        normal_format.space_before = pt(0)
        normal_format.space_after = pt(6)

        for style_name, space_before, space_after in (
            ("Heading 1", pt(0), pt(10)),
            ("Heading 2", pt(8), pt(6)),
        ):
            if style_name in document.styles:
                paragraph_format = document.styles[style_name].paragraph_format
                paragraph_format.line_spacing = 1.15
                paragraph_format.space_before = space_before
                paragraph_format.space_after = space_after

        for style_name in ("List Bullet", "List Number"):
            if style_name in document.styles:
                paragraph_format = document.styles[style_name].paragraph_format
                paragraph_format.line_spacing = 1.2
                paragraph_format.space_before = pt(0)
                paragraph_format.space_after = pt(3)

    def _add_paragraph(self, document, text: str, style: str | None = None):
        cleaned_text = self._sanitize_plain_text(text)
        paragraph = document.add_paragraph(style=style)
        if cleaned_text:
            paragraph.add_run(cleaned_text)
        else:
            paragraph.add_run("无")
        return paragraph

    def _sanitize_html(self, html_text: str) -> str:
        cleaned = self._sanitize_text_content(html_text)
        if "<" not in cleaned or ">" not in cleaned:
            return ""
        return cleaned.strip()

    def _sanitize_plain_text(self, text: str) -> str:
        cleaned = self._sanitize_text_content(text)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        return cleaned.strip()

    def _sanitize_text_content(self, text: str) -> str:
        cleaned = text.replace("\r\n", "\n").replace("\r", "\n").replace("\xa0", " ")
        cleaned = re.sub(r"\[Pasted ~\d+ lines\]", "", cleaned)
        filtered_lines = []
        for line in cleaned.split("\n"):
            stripped = line.strip()
            if not stripped:
                filtered_lines.append("")
                continue
            if _is_noise_text(stripped):
                continue
            filtered_lines.append(line)
        return "\n".join(filtered_lines)


class _DocxHTMLParser(HTMLParser):
    def __init__(self, document) -> None:
        super().__init__()
        self.document = document
        self.current_paragraph = None
        self.bold = False
        self.list_style = None
        self.ignored_tag_depth = 0

    def handle_starttag(self, tag, attrs):
        if tag in {"style", "script", "head", "title", "meta"}:
            self.ignored_tag_depth += 1
            return
        if self.ignored_tag_depth:
            return
        if tag in {"p", "div", "body", "html"}:
            return
        elif tag in {"strong", "b"}:
            self.bold = True
        elif tag == "br":
            if self.current_paragraph is not None and self.current_paragraph.text:
                self.current_paragraph.add_run().add_break()
        elif tag == "ul":
            self.list_style = "List Bullet"
        elif tag == "ol":
            self.list_style = "List Number"
        elif tag == "li":
            self.current_paragraph = self.document.add_paragraph(
                style=self.list_style or "List Bullet"
            )

    def handle_endtag(self, tag):
        if tag in {"style", "script", "head", "title", "meta"}:
            if self.ignored_tag_depth:
                self.ignored_tag_depth -= 1
            return
        if self.ignored_tag_depth:
            return
        if tag in {"strong", "b"}:
            self.bold = False
        elif tag in {"p", "div", "li"}:
            self.current_paragraph = None
        elif tag in {"ul", "ol"}:
            self.list_style = None

    def handle_data(self, data):
        if self.ignored_tag_depth:
            return
        text = data.replace("\xa0", " ")
        if not text.strip():
            return
        if _is_noise_text(text):
            return
        paragraph = self._ensure_paragraph()
        run = paragraph.add_run(text)
        run.bold = self.bold

    def _ensure_paragraph(self):
        if self.current_paragraph is None:
            self.current_paragraph = self.document.add_paragraph()
        return self.current_paragraph
